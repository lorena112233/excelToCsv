import pandas
from docx import Document

from globalVar import *

"""
NOTA: imported PANDAS to work with 
TABLES[column][row]
Usando Matrices
"""

#Separate the final text into diferent sections. Some will be similar in all of them, others will be variables 
introduction_DearText = ""         
loyalCustomers_welcomeB = ""    
firstTimeCustomers_welcome = ""       
end_kindRegardsText = ""
end_authorReference="" 
end_author = ""

# blank lines / between paragraphs
def agregarSaltosDeLinea(documento):
    for d in range(G_saltosDeLinea):
        documento.add_paragraph("")



def readFromExcel():
    print("Leyendo fichero / Reading file")

    # En este caso, no hay ni columna ni fila de encabezado, por lo que lo indico = NONE
    info = pandas.read_excel(G_file + '/welcomeLetters.xlsx', index_col=None, header=None)

    #print(info)
    crearDocx(info)


def crearDocx(tableInfo):

    loyalCustomers_welcomeB = tableInfo[3][0]
    firstTimeCustomers_welcome = tableInfo[4][0]
    CuerpoText = tableInfo[5][0]
    end_kindRegardsText = tableInfo[6][0]
    introduction_DearText = tableInfo[2][0]
    end_authorReference=tableInfo[7][0]
    end_author = tableInfo[8][0]

    # create new document
    document = Document()

    # Add paragraphs / custom. order
    for i in range(len(tableInfo.index)):

        document.add_paragraph(introduction_DearText + tableInfo[0][i])

        agregarSaltosDeLinea(document)

        if tableInfo[1][i] == "YES":
            document.add_paragraph(loyalCustomers_welcomeB + " " + CuerpoText)
        else:
            document.add_paragraph(firstTimeCustomers_welcome + " " + CuerpoText)

        agregarSaltosDeLinea(document)

        document.add_paragraph(end_kindRegardsText)

        agregarSaltosDeLinea(document)

        document.add_paragraph(end_author + "\n"+  end_authorReference)

        document.add_page_break()

    document.save(G_file + 'test.docx')

    print("Finished:", (len(tableInfo.index)), "Done")
 


if __name__ == "__main__":
    print("Starting app.")
    readFromExcel()
